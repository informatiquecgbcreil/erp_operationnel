import os
import shutil
import subprocess
from datetime import datetime, date

from docx import Document  # fallback

from app.models import SessionActivite, PresenceActivite, Participant, AtelierCapaciteMois

try:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Mm
except Exception:  # pragma: no cover
    DocxTemplate = None  # type: ignore
    InlineImage = None  # type: ignore
    Mm = None  # type: ignore


def _safe_filename(s: str) -> str:
    return "".join(c for c in s if c.isalnum() or c in ("-", "_", " ")).strip().replace(" ", "_")


def _archives_root(app) -> str:
    root = os.path.join(app.instance_path, "archives_emargements")
    os.makedirs(root, exist_ok=True)
    return root


def _month_folder(month: int) -> str:
    names = [
        "01_Janvier",
        "02_Fevrier",
        "03_Mars",
        "04_Avril",
        "05_Mai",
        "06_Juin",
        "07_Juillet",
        "08_Aout",
        "09_Septembre",
        "10_Octobre",
        "11_Novembre",
        "12_Decembre",
    ]
    if 1 <= month <= 12:
        return names[month - 1]
    return f"{month:02d}_Mois"


def _try_docx_to_pdf(docx_path: str) -> str | None:
    """Convert DOCX to PDF using LibreOffice (headless). Returns pdf path or None."""
    if not os.path.exists(docx_path):
        return None
    out_dir = os.path.dirname(docx_path)

    def _find_soffice() -> str | None:
        """Locate LibreOffice CLI binary.

        On Windows (especially when running as a service), LibreOffice is often
        installed but not available on PATH. We therefore:
        1) Check env LIBREOFFICE_PATH (user-configurable)
        2) Try shutil.which for common commands
        3) Probe common install locations on Windows
        """
        # 1) Explicit env override
        env_path = os.environ.get("LIBREOFFICE_PATH")
        if env_path and os.path.exists(env_path):
            return env_path

        # 2) PATH lookup
        for cmd in ("soffice.com", "soffice", "libreoffice"):
            p = shutil.which(cmd)
            if p:
                return p

        # 3) Common Windows locations
        if os.name == "nt":
            candidates = [
                r"C:\\Program Files\\LibreOffice\\program\\soffice.com",
                r"C:\\Program Files\\LibreOffice\\program\\soffice.exe",
                r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.com",
                r"C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",
            ]
            for c in candidates:
                if os.path.exists(c):
                    return c
        return None

    soffice = _find_soffice()
    if not soffice:
        return None

    # Dedicated headless profile (avoids AppData/locking issues under services)
    profile_dir = os.path.join(out_dir, "_lo_profile")
    os.makedirs(profile_dir, exist_ok=True)
    profile_uri = "file:///" + profile_dir.replace("\\", "/")
    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nolockcheck",
                "--norestore",
                f"--env:UserInstallation={profile_uri}",
                "--convert-to",
                "pdf",
                "--outdir",
                out_dir,
                docx_path,
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
        base = os.path.splitext(os.path.basename(docx_path))[0]
        pdf_path = os.path.join(out_dir, f"{base}.pdf")
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception:
        return None


def _install_default_templates(app) -> dict[str, str]:
    """Ensure Antoine's provided templates exist in instance/ and return their paths."""
    tpl_dir = os.path.join(app.instance_path, "docx_templates")
    os.makedirs(tpl_dir, exist_ok=True)

    # packaged assets (tracked in repo)
    assets_dir = os.path.join(os.path.dirname(__file__), "..", "assets")
    assets_dir = os.path.abspath(assets_dir)

    mapping = {
        "collectif": os.path.join(tpl_dir, "modele_collectif.docx"),
        "individuel": os.path.join(tpl_dir, "modele_individuel.docx"),
    }

    src_collectif = os.path.join(assets_dir, "modele_collectif.docx")
    src_indiv = os.path.join(assets_dir, "modele_individuel.docx")

    if os.path.exists(src_collectif) and not os.path.exists(mapping["collectif"]):
        shutil.copyfile(src_collectif, mapping["collectif"])
    if os.path.exists(src_indiv) and not os.path.exists(mapping["individuel"]):
        shutil.copyfile(src_indiv, mapping["individuel"])

    return mapping


def _format_date_fr(d: date | None) -> str:
    return d.strftime("%d/%m/%Y") if d else ""


def _inline_signature(app, signature_path: str | None):
    """Return an InlineImage for docxtpl, or empty string if missing."""
    if not signature_path or not os.path.exists(signature_path):
        return ""
    if DocxTemplate is None or InlineImage is None or Mm is None:
        return ""
    try:
        # 30mm wide is readable in a cell
        return InlineImage(DocxTemplate(""), signature_path, width=Mm(30))  # dummy template replaced in render
    except Exception:
        return ""


def _docxtpl_inline(template, signature_path: str | None):
    if not signature_path or not os.path.exists(signature_path):
        return ""
    if InlineImage is None or Mm is None:
        return ""
    try:
        return InlineImage(template, signature_path, width=Mm(30))
    except Exception:
        return ""


def generate_collectif_docx_pdf(app, atelier, session: SessionActivite):
    """Generate a DOCX (and try PDF) for a collective session.

    Uses docxtpl when a template is provided (Jinja in DOCX).
    Falls back to python-docx for simple templates.
    """
    dt = session.date_session or datetime.utcnow().date()
    y, m = dt.year, dt.month

    root = _archives_root(app)
    folder = os.path.join(root, _safe_filename(atelier.secteur), str(y), _safe_filename(atelier.nom), _month_folder(m))
    os.makedirs(folder, exist_ok=True)

    time_label = (session.heure_debut or "")
    if session.heure_fin:
        time_label = f"{time_label}-{session.heure_fin}" if time_label else session.heure_fin

    fname = f"{dt.isoformat()}__COLLECTIF__{_safe_filename(atelier.nom)}__{_safe_filename(time_label or 'session')}.docx"
    out_docx = os.path.join(folder, fname)

    defaults = _install_default_templates(app)
    template_path = atelier.modele_docx_collectif or defaults.get("collectif")

    presences = (
        PresenceActivite.query.filter_by(session_id=session.id)
        .join(Participant)
        .order_by(Participant.nom.asc(), Participant.prenom.asc())
        .all()
    )

    if DocxTemplate is not None and template_path and os.path.exists(template_path):
        tpl = DocxTemplate(template_path)
        participants = []
        for pr in presences:
            p = pr.participant
            participants.append(
                {
                    "nom": f"{(p.nom or '').upper()} {(p.prenom or '')}",
                    "email": p.email or "",
                    "ddn": _format_date_fr(p.date_naissance),
                    "sexe": p.genre or "",
                    "type": getattr(p, "type_public", None) or "H",
                    "ville": p.ville or "",
                    "signature": _docxtpl_inline(tpl, pr.signature_path),
                }
            )
        context = {
            "lieu": getattr(session, "lieu", None) or "",
            "date": dt.strftime("%d/%m/%Y"),
            "horaires": time_label,
            "titre": atelier.nom,
            "intervenant": "",
            "participants": participants,
        }
        tpl.render(context)
        tpl.save(out_docx)
    else:
        # Fallback: basic table using python-docx
        doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
        if not doc.paragraphs:
            doc.add_heading(f"Feuille d'émargement - {atelier.nom}", level=1)
        table = doc.tables[0] if doc.tables else doc.add_table(rows=1, cols=8)
        if len(table.rows) == 1:
            headers = ["Nom", "Email", "DDN", "Sexe", "Type", "Ville", "Motif", "Signature"]
            for i, h in enumerate(headers):
                table.cell(0, i).text = h
        while len(table.rows) > 1:
            tbl = table._tbl
            tbl.remove(table.rows[1]._tr)
        for pr in presences:
            p = pr.participant
            row = table.add_row().cells
            row[0].text = f"{(p.nom or '').upper()} {(p.prenom or '')}"
            row[1].text = p.email or ""
            row[2].text = _format_date_fr(p.date_naissance)
            row[3].text = p.genre or ""
            row[4].text = getattr(p, "type_public", None) or "H"
            row[5].text = p.ville or ""
            row[6].text = pr.motif or ""
            row[7].text = ""
        doc.save(out_docx)

    out_pdf = _try_docx_to_pdf(out_docx)
    return out_docx, out_pdf


def generate_individuel_mensuel_docx(app, atelier, annee: int, mois: int) -> str:
    """Generate a DOCX for an INDIVIDUEL_MENSUEL atelier for a month."""
    root = _archives_root(app)
    folder = os.path.join(root, _safe_filename(atelier.secteur), str(annee), _safe_filename(atelier.nom), _month_folder(mois))
    os.makedirs(folder, exist_ok=True)

    fname = f"{annee}-{mois:02d}__INDIVIDUEL__{_safe_filename(atelier.nom)}.docx"
    out_docx = os.path.join(folder, fname)

    defaults = _install_default_templates(app)
    template_path = atelier.modele_docx_individuel or defaults.get("individuel")

    # Fetch RDV sessions
    sessions = (
        SessionActivite.query.filter_by(atelier_id=atelier.id, session_type="INDIVIDUEL_MENSUEL", is_deleted=False)
        .filter(SessionActivite.rdv_date.isnot(None))
        .filter(SessionActivite.rdv_date >= date(annee, mois, 1))
        .all()
    )

    # Only keep month
    sessions = [s for s in sessions if s.rdv_date and s.rdv_date.year == annee and s.rdv_date.month == mois]
    sessions.sort(key=lambda s: (s.rdv_date or date(annee, mois, 1), s.rdv_debut or ""))

    # Build rows (one line per RDV). We take the 1st presence as the participant for individual.
    rows = []
    for s in sessions:
        pr = (
            PresenceActivite.query.filter_by(session_id=s.id)
            .join(Participant)
            .order_by(Participant.nom.asc(), Participant.prenom.asc())
            .first()
        )
        if not pr:
            continue
        p = pr.participant
        heures = ""
        if s.rdv_debut and s.rdv_fin:
            heures = f"{s.rdv_debut} - {s.rdv_fin}"
        elif s.rdv_debut:
            heures = s.rdv_debut
        motif = pr.motif or ""
        if pr.motif_autre:
            motif = f"{motif} / {pr.motif_autre}" if motif else pr.motif_autre

        rows.append(
            {
                "nom": f"{(p.nom or '').upper()} {(p.prenom or '')}",
                "email": p.email or "",
                "ddn": _format_date_fr(p.date_naissance),
                "sexe": p.genre or "",
                "type": getattr(p, "type_public", None) or "H",
                "da": _format_date_fr(s.rdv_date),  # IMPORTANT: your template currently uses p.da
                "heures": heures,
                "motif": motif,
                "ville": p.ville or "",
                "signature": None,  # filled below if docxtpl
                "_sig_path": pr.signature_path,
            }
        )

    if DocxTemplate is not None and template_path and os.path.exists(template_path):
        tpl = DocxTemplate(template_path)
        for r in rows:
            r["signature"] = _docxtpl_inline(tpl, r.pop("_sig_path", None))
        context = {
            "lieu": getattr(atelier, "lieu", None) or "",
            "mois": mois,
            "annee": annee,
            "titre": atelier.nom,
            "intervenant": "",
            "participants": rows,
        }
        tpl.render(context)
        tpl.save(out_docx)
    else:
        # fallback
        doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
        doc.add_paragraph(f"{atelier.nom} - {mois:02d}/{annee}")
        table = doc.tables[0] if doc.tables else doc.add_table(rows=1, cols=9)
        if len(table.rows) == 1:
            headers = ["Nom", "Email", "DDN", "Sexe", "Type", "Date", "Heures", "Motif", "Ville"]
            for i, h in enumerate(headers):
                table.cell(0, i).text = h
        while len(table.rows) > 1:
            tbl = table._tbl
            tbl.remove(table.rows[1]._tr)
        for r in rows:
            row = table.add_row().cells
            row[0].text = r.get("nom", "")
            row[1].text = r.get("email", "")
            row[2].text = r.get("ddn", "")
            row[3].text = r.get("sexe", "")
            row[4].text = r.get("type", "H")
            row[5].text = r.get("da", "")
            row[6].text = r.get("heures", "")
            row[7].text = r.get("motif", "")
            row[8].text = r.get("ville", "")
        doc.save(out_docx)

    return out_docx


def finalize_individuel_mensuel_pdf(app, atelier, annee: int, mois: int) -> str | None:
    docx_path = generate_individuel_mensuel_docx(app, atelier, annee, mois)
    return _try_docx_to_pdf(docx_path)


def generate_participant_bilan_docx(app, participant, rows: list[dict]) -> str:
    folder = os.path.join(app.instance_path, "archives_pedagogie")
    os.makedirs(folder, exist_ok=True)
    fname = f"bilan_{participant.id}_{_safe_filename(participant.nom)}_{_safe_filename(participant.prenom)}.docx"
    out_docx = os.path.join(folder, fname)

    template_path = os.path.join(app.instance_path, "docx_templates", "bilan_pedagogique.docx")

    if DocxTemplate and os.path.exists(template_path):
        tpl = DocxTemplate(template_path)
        context = {
            "participant": {
                "nom": participant.nom,
                "prenom": participant.prenom,
                "email": participant.email or "",
                "ville": participant.ville or "",
            },
            "rows": rows,
            "date": date.today().strftime("%d/%m/%Y"),
        }
        tpl.render(context)
        tpl.save(out_docx)
    else:
        doc = Document(template_path) if template_path and os.path.exists(template_path) else Document()
        doc.add_heading("Bilan pédagogique", level=1)
        doc.add_paragraph(f"Participant : {participant.nom} {participant.prenom}")
        if participant.email:
            doc.add_paragraph(f"Email : {participant.email}")
        if participant.ville:
            doc.add_paragraph(f"Ville : {participant.ville}")
        doc.add_paragraph(f"Date : {date.today().strftime('%d/%m/%Y')}")
        table = doc.add_table(rows=1, cols=4)
        headers = ["Référentiel", "Compétence", "Date", "Atelier"]
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
        for r in rows:
            row = table.add_row().cells
            row[0].text = r.get("referentiel", "")
            row[1].text = r.get("competence", "")
            row[2].text = r.get("date", "")
            row[3].text = r.get("atelier", "")
        doc.save(out_docx)

    return out_docx


def generate_participant_bilan_pdf(app, participant, rows: list[dict]) -> str | None:
    docx_path = generate_participant_bilan_docx(app, participant, rows)
    return _try_docx_to_pdf(docx_path)
